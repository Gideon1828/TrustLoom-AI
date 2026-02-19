"""
Resume Text Processing Pipeline
Extracts and cleans text from PDF and DOCX resume files
Author: Freelancer Trust Evaluation Team
Version: 1.0
"""

import re
import logging
from pathlib import Path
from typing import Optional, Union
import PyPDF2
import docx
from config.config import FileProcessingConfig

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Parser for extracting and cleaning text from resume files (PDF and DOCX)
    """
    
    def __init__(self):
        self.max_length = FileProcessingConfig.MAX_RESUME_LENGTH
        self.min_length = FileProcessingConfig.MIN_RESUME_LENGTH
        self.text_cleaning_enabled = FileProcessingConfig.TEXT_CLEANING_ENABLED
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract raw text from resume file (PDF or DOCX)
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted raw text as string
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                raw_text = self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                raw_text = self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx, .doc")
            
            logger.info(f"Successfully extracted {len(raw_text)} characters from {file_path.name}")
            return raw_text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        text = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def clean_text(self, raw_text: str) -> str:
        """
        Clean extracted text by removing formatting, special characters, and excessive whitespace
        
        Args:
            raw_text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not self.text_cleaning_enabled:
            return raw_text
        
        text = raw_text
        
        # Remove URLs
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        
        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove special characters but preserve alphanumeric, spaces, and basic punctuation
        text = re.sub(r'[^\w\s\.,;:\-\(\)\[\]\{\}\'\"\/\+\#\&]', ' ', text)
        
        # Remove extra bullet points and formatting characters
        text = re.sub(r'[•●○▪▫■□★☆▸▹►▻]', '', text)
        
        # Remove multiple consecutive dots (like .........)
        text = re.sub(r'\.{2,}', '.', text)
        
        # Remove multiple consecutive dashes or underscores
        text = re.sub(r'-{2,}', '-', text)
        text = re.sub(r'_{2,}', '_', text)
        
        # Remove form field remnants (like ________)
        text = re.sub(r'_{3,}', '', text)
        
        # Replace multiple SPACES with single space (but preserve newlines!)
        # Split by lines to preserve line structure
        lines = text.split('\n')
        lines = [re.sub(r'[ \t]+', ' ', line) for line in lines]  # Only collapse spaces/tabs, not newlines
        text = '\n'.join(lines)
        
        # Replace multiple newlines with double newline (preserve some spacing between sections)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove empty lines
        lines = [line for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        
        # Final trim
        text = text.strip()
        
        logger.info(f"Text cleaned: {len(raw_text)} -> {len(text)} characters")
        return text
    
    def process_resume(self, file_path: Union[str, Path]) -> str:
        """
        Complete pipeline: Extract and clean text from resume
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Clean, processed text ready for NLP analysis
            
        Raises:
            ValueError: If text is too short or too long
        """
        # Extract raw text
        raw_text = self.extract_text(file_path)
        
        # Clean text
        clean_text = self.clean_text(raw_text)
        
        # Validate text length
        text_length = len(clean_text)
        
        if text_length < self.min_length:
            raise ValueError(
                f"Resume text too short ({text_length} chars). "
                f"Minimum required: {self.min_length} chars"
            )
        
        if text_length > self.max_length:
            logger.warning(
                f"Resume text exceeds maximum length ({text_length} chars). "
                f"Truncating to {self.max_length} chars"
            )
            clean_text = clean_text[:self.max_length]
        
        logger.info(f"Resume processed successfully: {text_length} characters")
        return clean_text


# Convenience functions for direct use
def extract_text_from_resume(file_path: Union[str, Path]) -> str:
    """
    Convenience function to extract raw text from resume
    
    Args:
        file_path: Path to resume file
        
    Returns:
        Raw extracted text
    """
    parser = ResumeParser()
    return parser.extract_text(file_path)


def clean_text(raw_text: str) -> str:
    """
    Convenience function to clean text
    
    Args:
        raw_text: Raw text to clean
        
    Returns:
        Cleaned text
    """
    parser = ResumeParser()
    return parser.clean_text(raw_text)


def process_resume(file_path: Union[str, Path]) -> str:
    """
    Convenience function for complete resume processing pipeline
    
    Args:
        file_path: Path to resume file
        
    Returns:
        Processed, clean text ready for analysis
    """
    parser = ResumeParser()
    return parser.process_resume(file_path)


# Example usage and testing
if __name__ == "__main__":
    import sys
    
    print("=" * 60)
    print("RESUME TEXT PROCESSING PIPELINE - TEST")
    print("=" * 60)
    
    if len(sys.argv) > 1:
        # Test with provided file
        test_file = sys.argv[1]
        print(f"\nProcessing: {test_file}")
        
        try:
            parser = ResumeParser()
            
            # Extract raw text
            print("\n--- Extracting raw text ---")
            raw = parser.extract_text(test_file)
            print(f"Raw text length: {len(raw)} characters")
            print(f"First 200 chars:\n{raw[:200]}...")
            
            # Clean text
            print("\n--- Cleaning text ---")
            cleaned = parser.clean_text(raw)
            print(f"Cleaned text length: {len(cleaned)} characters")
            print(f"First 200 chars:\n{cleaned[:200]}...")
            
            # Full processing
            print("\n--- Full processing pipeline ---")
            processed = parser.process_resume(test_file)
            print(f"Final text length: {len(processed)} characters")
            print("✓ Processing successful!")
            
        except Exception as e:
            print(f"✗ Error: {str(e)}")
            sys.exit(1)
    else:
        print("\nUsage: python resume_parser.py <path_to_resume>")
        print("Example: python resume_parser.py ../data/sample_resumes/resume.pdf")
    
    print("\n" + "=" * 60)
